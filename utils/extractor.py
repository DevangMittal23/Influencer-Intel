import io
import re
import httpx
import pdfplumber
from docx import Document
from newspaper import Article
from langdetect import detect, LangDetectException


def _detect_language(text: str) -> str:
    try:
        sample = text[:500]
        lang = detect(sample)
        # Classify as mixed if we detect both latin and devanagari scripts
        has_devanagari = bool(re.search(r'[\u0900-\u097F]', sample))
        has_latin = bool(re.search(r'[a-zA-Z]{3,}', sample))
        if has_devanagari and has_latin:
            return "mixed"
        return lang
    except LangDetectException:
        return "unknown"


def _strip_html(html: str) -> str:
    text = re.sub(r'<[^>]+>', ' ', html)
    return re.sub(r'\s+', ' ', text).strip()


def extract_text(source, source_type: str, source_name: str = "") -> dict:
    result = {
        "text": "",
        "source_name": source_name,
        "source_type": source_type,
        "char_count": 0,
        "language": "unknown",
        "extraction_success": False,
        "error": None,
    }

    try:
        if source_type == "pdf":
            with pdfplumber.open(io.BytesIO(source) if isinstance(source, bytes) else source) as pdf:
                pages = [page.extract_text() or "" for page in pdf.pages]
            result["text"] = "\n".join(pages).strip()

        elif source_type == "docx":
            doc = Document(io.BytesIO(source) if isinstance(source, bytes) else source)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            table_cells = [
                cell.text for table in doc.tables
                for row in table.rows for cell in row.cells if cell.text.strip()
            ]
            result["text"] = "\n".join(paragraphs + table_cells).strip()

        elif source_type == "url":
            try:
                article = Article(source)
                article.download()
                article.parse()
                title = article.title or ""
                body = article.text or ""
                result["text"] = f"{title}\n\n{body}".strip() if title else body.strip()
            except Exception:
                resp = httpx.get(source, timeout=15, follow_redirects=True)
                result["text"] = _strip_html(resp.text)

        elif source_type in ("txt", "raw_text"):
            text = source.decode("utf-8", errors="replace") if isinstance(source, bytes) else source
            result["text"] = re.sub(r'\s+', ' ', text).strip()

        if not result["text"]:
            raise ValueError("No text could be extracted.")

        result["extraction_success"] = True
        result["char_count"] = len(result["text"])
        result["language"] = _detect_language(result["text"])

    except Exception as e:
        result["error"] = str(e)
        result["extraction_success"] = False

    return result
